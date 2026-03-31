#!/usr/bin/env python3
"""
Tender Document Generator - CLI версия
Для соответствия ТЗ - автоматическое заполнение без GUI
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from docx import Document


class TenderDocumentGenerator:
    """Генератор тендерных документов"""
    
    def __init__(self, data_dir: str, templates_dir: str, output_dir: str):
        self.data_dir = Path(data_dir)
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.current_date = datetime.now()
        
        # Загрузка данных
        self.profile = self._load_json("profile.json")
        self.tender = self._load_json("tender.json")
        self.calc = self._load_json("calc.json")
    
    def _load_json(self, filename: str) -> dict:
        """Загрузить JSON файл"""
        with open(self.data_dir / filename, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def _format_date_ru(self) -> str:
        """Форматировать дату"""
        months = ['', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
        return f'«{self.current_date.day:02d}» {months[self.current_date.month]} {self.current_date.year} года'
    
    def generate_all(self):
        """Сгенерировать все документы"""
        print("="*60)
        print("ГЕНЕРАЦИЯ ТЕНДЕРНЫХ ДОКУМЕНТОВ")
        print("="*60)
        
        # Анкета
        print("\n[1/3] Генерация Анкеты участника...")
        self._generate_anketa()
        print("   ✓ Анкета сгенерирована")
        
        # Заявка
        print("\n[2/3] Генерация Заявки на участие...")
        self._generate_zayavka()
        print("   ✓ Заявка сгенерирована")
        
        # Предложение о цене
        print("\n[3/3] Генерация Предложения о цене...")
        self._generate_predlozhenie()
        print("   ✓ Предложение о цене сгенерировано")
        
        print("\n" + "="*60)
        print(f"ГОТОВО! Документы сохранены в: {self.output_dir}")
        print("="*60)
    
    def _generate_anketa(self):
        """Анкета участника"""
        template = self.templates_dir / "01_Анкета_участника_шаблон.docx"
        output = self.output_dir / "01_Анкета_участника_заполненная.docx"
        
        doc = Document(template)
        profile = self.profile
        
        # Заголовок
        for para in doc.paragraphs:
            if '[Краткое наименование участника] [ИНН]' in para.text:
                para.text = f"{profile['company']['short_name']} {profile['company']['inn']}"
        
        # Таблица
        if doc.tables:
            table = doc.tables[0]
            
            mappings = [
                (1, profile['company']['full_name']),
                (2, profile['company']['short_name']),
                (3, profile['company']['inn']),
                (4, profile['company']['kpp']),
                (5, profile['company']['ogrn']),
                (6, profile['company']['legal_address_full']),
            ]
            
            for row_idx, value in mappings:
                if row_idx < len(table.rows):
                    table.rows[row_idx].cells[1].text = str(value)
            
            # Банк
            if 7 < len(table.rows):
                bank = profile['bank']
                text = f"{bank['account']}\n{bank['name']}\nк/с {bank['correspondent_account']}\nБИК {bank['bik']}"
                table.rows[7].cells[1].text = text
            
            # Контакты
            if 8 < len(table.rows):
                contact = profile['contact']
                text = f"{contact['responsible_name_full']}\n{contact['email']}\n{contact['phone']}"
                table.rows[8].cells[1].text = text
            
            # Подписант
            if len(doc.tables) > 1:
                sign = doc.tables[1]
                if sign.rows:
                    cells = sign.rows[0].cells
                    cells[0].text = profile['signatory']['position']
                    if len(cells) >= 3:
                        cells[2].text = profile['signatory']['name_short']
        
        doc.save(output)
    
    def _generate_zayavka(self):
        """Заявка на участие"""
        template = self.templates_dir / "02_Заявка_на_участие_в_закупке_шаблон.docx"
        output = self.output_dir / "02_Заявка_на_участие_в_закупке_заполненная.docx"
        
        doc = Document(template)
        profile = self.profile
        tender = self.tender
        
        date_str = self._format_date_ru()
        outgoing = f"№{self.current_date.strftime('%d%m')}/1"
        
        for para in doc.paragraphs:
            text = para.text
            
            if '[Дата в формате «26» марта 2026 года]' in text:
                text = text.replace('[Дата в формате «26» марта 2026 года]', date_str)
            
            if '[Исх. номер заявки]' in text:
                text = text.replace('[Исх. номер заявки]', outgoing)
            
            if '[Полное наименование участника]' in text:
                text = text.replace('[Полное наименование участника]', profile['company']['full_name'])
            
            if '[Юридический адрес]' in text:
                text = text.replace('[Юридический адрес]', profile['company']['legal_address_full'])
            
            if '[Предмет закупки]' in text:
                text = text.replace('[Предмет закупки]', tender['subject'])
            
            if '[Срок действия предложения, дней]' in text:
                text = text.replace('[Срок действия предложения, дней]', str(tender['offer_validity_days']))
            
            para.text = text
        
        # Подписант
        if doc.tables:
            sign = doc.tables[0]
            if sign.rows:
                cells = sign.rows[0].cells
                cells[0].text = profile['signatory']['position']
                if len(cells) >= 3:
                    cells[2].text = profile['signatory']['name_short']
        
        doc.save(output)
    
    def _generate_predlozhenie(self):
        """Предложение о цене"""
        template = self.templates_dir / "03_Предложение_о_цене_договора_шаблон.docx"
        output = self.output_dir / "03_Предложение_о_цене_договора_заполненное.docx"
        
        doc = Document(template)
        profile = self.profile
        calc = self.calc
        
        # Заголовок
        for para in doc.paragraphs:
            if '[Краткое наименование участника] [ИНН]' in para.text:
                para.text = f"{profile['company']['short_name']} {profile['company']['inn']}"
        
        if not doc.tables:
            doc.save(output)
            return
        
        # Позиции
        items_table = doc.tables[0]
        items = calc['items']
        
        for idx, item in enumerate(items):
            row_idx = idx + 1
            while len(items_table.rows) <= row_idx:
                items_table.add_row()
            
            row = items_table.rows[row_idx]
            if len(row.cells) >= 7:
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = item['quote_name']
                row.cells[2].text = item['offer_unit']
                row.cells[3].text = f"{item['unit_price_wo_vat']:.2f}"
                row.cells[4].text = f"{item['unit_price_with_delivery_wo_vat']:.2f}"
                row.cells[5].text = str(item['offer_qty'])
                row.cells[6].text = f"{item['line_total_wo_vat']:,.2f}"
        
        # Итоги
        if len(doc.tables) > 1:
            totals = doc.tables[1]
            if len(totals.rows) >= 3:
                totals.rows[0].cells[1].text = f"{calc['subtotal_wo_vat']:,.2f}"
                totals.rows[1].cells[1].text = f"{calc['vat_amount']:,.2f}"
                totals.rows[2].cells[1].text = f"{calc['total_with_vat']:,.2f}"
        
        # Подписант
        if len(doc.tables) > 2:
            sign = doc.tables[2]
            if sign.rows:
                cells = sign.rows[0].cells
                cells[0].text = profile['signatory']['position']
                if len(cells) >= 3:
                    cells[2].text = profile['signatory']['name_short']
        
        doc.save(output)


def main():
    """Точка входа CLI"""
    if len(sys.argv) < 4:
        print("Использование:")
        print("  python generator_cli.py <data_dir> <templates_dir> <output_dir>")
        print("\nПример:")
        print("  python generator_cli.py ./data ./templates ./output")
        sys.exit(1)
    
    data_dir = sys.argv[1]
    templates_dir = sys.argv[2]
    output_dir = sys.argv[3]
    
    generator = TenderDocumentGenerator(data_dir, templates_dir, output_dir)
    generator.generate_all()


if __name__ == "__main__":
    main()
