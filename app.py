"""
Tender Doc Generator GUI v4.0
Исправлены баги с отображением файлов и генерацией всех документов
"""

import os
import sys
import threading
import json
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, StringVar

try:
    import customtkinter as ctk
    from docx import Document
    CTK_AVAILABLE = True
except ImportError as e:
    print(f"Import error: {e}")
    CTK_AVAILABLE = False


IOS_BLUE = "#007AFF"
IOS_GREEN = "#34C759"
IOS_RED = "#FF3B30"
IOS_GRAY = "#8E8E93"
IOS_LIGHT_GRAY = "#F2F2F7"
IOS_WHITE = "#FFFFFF"
IOS_BLACK = "#000000"


class DocumentGenerator:
    """Генератор документов"""
    
    def __init__(self, templates_dir, output_dir, progress_callback=None):
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.progress_callback = progress_callback
        self.current_date = datetime.now()
    
    def _format_date_ru(self):
        months = ['', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
        return f'«{self.current_date.day:02d}» {months[self.current_date.month]} {self.current_date.year} года'
    
    def generate_all(self, profile, tender, calc):
        results = []
        
        docs = [
            ('anketa', '01_Анкета_участника_шаблон.docx', '01_Анкета_участника_заполненная.docx'),
            ('zayavka', '02_Заявка_на_участие_в_закупке_шаблон.docx', '02_Заявка_на_участие_в_закупке_заполненная.docx'),
            ('predlozhenie', '03_Предложение_о_цене_договора_шаблон.docx', '03_Предложение_о_цене_договора_заполненное.docx')
        ]
        
        for idx, (doc_id, template, output) in enumerate(docs, 1):
            if self.progress_callback:
                self.progress_callback(f"Генерация документа {idx} из 3...", idx / 3 * 100)
            
            try:
                method = getattr(self, f'_generate_{doc_id}')
                output_path = method(template, output, profile, tender, calc)
                results.append({'status': 'success', 'output': str(output_path)})
            except Exception as e:
                results.append({'status': 'error', 'error': str(e)})
        
        return results
    
    def _generate_anketa(self, template, output, profile, tender, calc):
        """Анкета участника"""
        template_path = self.templates_dir / template
        output_path = self.output_dir / output
        
        doc = Document(template_path)
        
        # Заголовок
        for para in doc.paragraphs:
            if '[Краткое наименование участника] [ИНН]' in para.text:
                para.text = para.text.replace(
                    '[Краткое наименование участника] [ИНН]',
                    f"{profile['company']['short_name']} {profile['company']['inn']}"
                )
        
        # Таблица
        if doc.tables:
            table = doc.tables[0]
            
            field_map = [
                (1, profile['company']['full_name']),
                (2, profile['company']['short_name']),
                (3, profile['company']['inn']),
                (4, profile['company']['kpp']),
                (5, profile['company']['ogrn']),
                (6, profile['company']['legal_address_full']),
            ]
            
            for row_idx, value in field_map:
                if row_idx < len(table.rows):
                    table.rows[row_idx].cells[1].text = str(value)
            
            # Банк
            if 7 < len(table.rows):
                bank = profile['bank']
                text = f"{bank['account']}\n{bank['name']}\nк/с {bank['correspondent_account']}\nБИК {bank['bik']}"
                table.rows[7].cells[1].text = text
            
            # Контакты
            if 8 < len(table.rows):
                contact = profile['contact']
                text = f"{contact['responsible_name_full']}\n{contact['email']}\n{contact['phone']}"
                table.rows[8].cells[1].text = text
            
            # Подписант
            if len(doc.tables) > 1:
                sign = doc.tables[1]
                if sign.rows:
                    cells = sign.rows[0].cells
                    cells[0].text = profile['signatory']['position']
                    if len(cells) >= 3:
                        cells[2].text = profile['signatory']['name_short']
        
        doc.save(output_path)
        return output_path
    
    def _generate_zayavka(self, template, output, profile, tender, calc):
        """Заявка на участие - ИСПРАВЛЕННАЯ"""
        template_path = self.templates_dir / template
        output_path = self.output_dir / output
        
        doc = Document(template_path)
        
        date_str = self._format_date_ru()
        outgoing = f"№{self.current_date.strftime('%d%m')}/1"
        
        # Замена во всех параграфах
        for para in doc.paragraphs:
            text = para.text
            
            if '[Дата в формате «26» марта 2026 года]' in text:
                text = text.replace('[Дата в формате «26» марта 2026 года]', date_str)
            if '[Дата в формате' in text and '2026 года]' in text:
                text = text.replace(text[text.find('[Дата в формате'):text.find('2026 года]')+10], date_str)
            
            if '[Исх. номер заявки]' in text:
                text = text.replace('[Исх. номер заявки]', outgoing)
            
            if '[Полное наименование участника]' in text:
                text = text.replace('[Полное наименование участника]', profile['company']['full_name'])
            
            if '[Юридический адрес]' in text:
                text = text.replace('[Юридический адрес]', profile['company']['legal_address_full'])
            
            if '[Предмет закупки]' in text:
                text = text.replace('[Предмет закупки]', tender.get('subject', ''))
            
            if '[Срок действия предложения, дней]' in text:
                text = text.replace('[Срок действия предложения, дней]', str(tender.get('offer_validity_days', '')))
            
            para.text = text
        
        # Подписант в таблице
        if doc.tables:
            sign = doc.tables[0]
            if sign.rows:
                cells = sign.rows[0].cells
                cells[0].text = profile['signatory']['position']
                if len(cells) >= 3:
                    cells[2].text = profile['signatory']['name_short']
        
        doc.save(output_path)
        return output_path
    
    def _generate_predlozhenie(self, template, output, profile, tender, calc):
        """Предложение о цене"""
        template_path = self.templates_dir / template
        output_path = self.output_dir / output
        
        doc = Document(template_path)
        
        # Заголовок
        for para in doc.paragraphs:
            if '[Краткое наименование участника] [ИНН]' in para.text:
                para.text = para.text.replace(
                    '[Краткое наименование участника] [ИНН]',
                    f"{profile['company']['short_name']} {profile['company']['inn']}"
                )
        
        if not doc.tables:
            doc.save(output_path)
            return output_path
        
        # Таблица позиций
        items_table = doc.tables[0]
        items = calc.get('items', [])
        
        for idx, item in enumerate(items):
            row_idx = idx + 1
            while len(items_table.rows) <= row_idx:
                items_table.add_row()
            
            row = items_table.rows[row_idx]
            if len(row.cells) >= 7:
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = item.get('quote_name', '')
                row.cells[2].text = item.get('offer_unit', '')
                row.cells[3].text = f"{item.get('unit_price_wo_vat', 0):.2f}"
                row.cells[4].text = f"{item.get('unit_price_with_delivery_wo_vat', 0):.2f}"
                row.cells[5].text = str(item.get('offer_qty', ''))
                row.cells[6].text = f"{item.get('line_total_wo_vat', 0):,.2f}"
        
        # Итоги
        if len(doc.tables) > 1:
            totals = doc.tables[1]
            if len(totals.rows) >= 3:
                totals.rows[0].cells[1].text = f"{calc.get('subtotal_wo_vat', 0):,.2f}"
                totals.rows[1].cells[1].text = f"{calc.get('vat_amount', 0):,.2f}"
                totals.rows[2].cells[1].text = f"{calc.get('total_with_vat', 0):,.2f}"
        
        # Подписант
        if len(doc.tables) > 2:
            sign = doc.tables[2]
            if sign.rows:
                cells = sign.rows[0].cells
                cells[0].text = profile['signatory']['position']
                if len(cells) >= 3:
                    cells[2].text = profile['signatory']['name_short']
        
        doc.save(output_path)
        return output_path


class TenderGeneratorApp:
    """Главное окно"""
    
    def __init__(self):
        self.root = ctk.CTk()
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        
        self.root.title("Tender Doc Generator v4.0")
        self.root.geometry("500x650")
        self.root.configure(fg_color=IOS_LIGHT_GRAY)
        self.root.resizable(False, False)
        
        # Центрирование
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        self.root.geometry(f"+{(sw-500)//2}+{(sh-650)//2}")
        
        self.output_dir = None
        self.template_dir = None
        
        self._create_ui()
    
    def _create_ui(self):
        main = ctk.CTkFrame(self.root, fg_color=IOS_LIGHT_GRAY)
        main.pack(fill="both", expand=True, padx=12, pady=12)
        
        # Header
        header = ctk.CTkFrame(main, fg_color=IOS_WHITE, corner_radius=10, height=70)
        header.pack(fill="x", pady=(0, 10))
        header.pack_propagate(False)
        
        ctk.CTkLabel(header, text="Tender Doc Generator", 
                    font=("SF Pro Display", 18, "bold")).pack(pady=(8, 0))
        ctk.CTkLabel(header, text="Автозаполнение тендерных документов", 
                    font=("SF Pro Display", 11), text_color=IOS_GRAY).pack()
        
        # Data card
        data_card = self._create_card(main, "ШАГ 1: ДАННЫЕ")
        
        # Profile - ОТДЕЛЬНАЯ переменная для каждого
        self.profile_var = StringVar(value="Не выбрано")
        self.profile_frame = self._create_file_selector(data_card, "profile.json", self.profile_var, "profile")
        self.profile_frame.pack(fill="x", padx=10, pady=2)
        
        # Tender - ОТДЕЛЬНАЯ переменная
        self.tender_var = StringVar(value="Не выбрано")
        self.tender_frame = self._create_file_selector(data_card, "tender.json", self.tender_var, "tender")
        self.tender_frame.pack(fill="x", padx=10, pady=2)
        
        # Calc - ОТДЕЛЬНАЯ переменная
        self.calc_var = StringVar(value="Не выбрано")
        self.calc_frame = self._create_file_selector(data_card, "calc.json", self.calc_var, "calc")
        self.calc_frame.pack(fill="x", padx=10, pady=(2, 8))
        
        # Templates card
        tmpl_card = self._create_card(main, "ШАГ 2: ШАБЛОНЫ")
        
        self.tmpl_btn = ctk.CTkButton(tmpl_card, text="Выбрать папку с шаблонами",
                                     height=36, corner_radius=8, fg_color=IOS_BLUE,
                                     font=("SF Pro Display", 13), command=self._select_templates)
        self.tmpl_btn.pack(fill="x", padx=10, pady=8)
        
        self.tmpl_label = ctk.CTkLabel(tmpl_card, text="Шаблоны не выбраны",
                                      font=("SF Pro Display", 11), text_color=IOS_GRAY)
        self.tmpl_label.pack(padx=10, pady=(0, 8))
        
        # Output card
        out_card = self._create_card(main, "ШАГ 3: СОХРАНЕНИЕ")
        
        self.out_btn = ctk.CTkButton(out_card, text="Выбрать папку для сохранения",
                                    height=36, corner_radius=8, fg_color=IOS_BLUE,
                                    font=("SF Pro Display", 13), command=self._select_output)
        self.out_btn.pack(fill="x", padx=10, pady=8)
        
        self.out_label = ctk.CTkLabel(out_card, text="Папка не выбрана",
                                     font=("SF Pro Display", 11), text_color=IOS_GRAY)
        self.out_label.pack(padx=10, pady=(0, 8))
        
        # Progress
        self.progress_frame = ctk.CTkFrame(main, fg_color="transparent")
        self.progress_bar = ctk.CTkProgressBar(self.progress_frame, mode="determinate",
                                               progress_color=IOS_BLUE, height=6)
        self.progress_bar.pack(fill="x", padx=10)
        self.progress_bar.set(0)
        
        self.status_label = ctk.CTkLabel(self.progress_frame, text="",
                                        font=("SF Pro Display", 11), text_color=IOS_GRAY)
        self.status_label.pack(pady=(4, 0))
        
        # Generate button
        self.gen_btn = ctk.CTkButton(main, text="▶  СГЕНЕРИРОВАТЬ ДОКУМЕНТЫ",
                                    height=45, corner_radius=10, fg_color=IOS_GREEN,
                                    font=("SF Pro Display", 14, "bold"), command=self._generate)
        self.gen_btn.pack(fill="x", pady=(0, 6))
        
        # Footer
        ctk.CTkLabel(main, text="Tender Doc Generator v4.0 • 2025",
                    font=("SF Pro Display", 10), text_color=IOS_GRAY).pack()
    
    def _create_card(self, parent, title):
        card = ctk.CTkFrame(parent, fg_color=IOS_WHITE, corner_radius=10)
        card.pack(fill="x", pady=(0, 8))
        ctk.CTkLabel(card, text=title, font=("SF Pro Display", 12, "bold"),
                    text_color=IOS_GRAY).pack(anchor="w", padx=10, pady=(8, 4))
        return card
    
    def _create_file_selector(self, parent, label, var, attr_name):
        """Создать селектор файла с отдельной переменной"""
        frame = ctk.CTkFrame(parent, fg_color="transparent")
        
        ctk.CTkLabel(frame, text=label, font=("SF Pro Display", 11),
                    text_color=IOS_GRAY).pack(anchor="w")
        
        inner = ctk.CTkFrame(frame, fg_color=IOS_LIGHT_GRAY, corner_radius=6)
        inner.pack(fill="x", pady=(2, 0))
        
        lbl = ctk.CTkLabel(inner, textvariable=var, font=("SF Pro Display", 12),
                          text_color=IOS_GRAY, anchor="w")
        lbl.pack(side="left", padx=8, fill="x", expand=True)
        
        def select():
            path = filedialog.askopenfilename(filetypes=[("JSON", "*.json")])
            if path:
                var.set(Path(path).name)
                setattr(self, f"{attr_name}_path", Path(path))
                lbl.configure(text_color=IOS_BLACK)
        
        ctk.CTkButton(inner, text="Выбрать", width=70, height=24, corner_radius=6,
                     fg_color=IOS_BLUE, font=("SF Pro Display", 11), command=select).pack(side="right", padx=4, pady=4)
        
        setattr(self, f"{attr_name}_path", None)
        return frame
    
    def _select_templates(self):
        path = filedialog.askdirectory(title="Выберите папку с шаблонами")
        if path:
            self.template_dir = Path(path)
            self.tmpl_label.configure(text=f"Выбрано: {self.template_dir}", text_color=IOS_BLACK)
    
    def _select_output(self):
        path = filedialog.askdirectory(title="Выберите папку для сохранения")
        if path:
            self.output_dir = Path(path)
            display = str(self.output_dir)
            if len(display) > 40:
                display = display[:18] + "..." + display[-19:]
            self.out_label.configure(text=f"Сохранение в: {display}", text_color=IOS_BLACK)
    
    def _progress_callback(self, msg, pct):
        self.status_label.configure(text=msg)
        self.progress_bar.set(pct / 100)
        self.root.update_idletasks()
    
    def _generate(self):
        errors = []
        
        profile = getattr(self, 'profile_path', None)
        tender = getattr(self, 'tender_path', None)
        calc = getattr(self, 'calc_path', None)
        
        if not profile: errors.append("Не выбран profile.json")
        if not tender: errors.append("Не выбран tender.json")
        if not calc: errors.append("Не выбран calc.json")
        if not self.template_dir: errors.append("Не выбрана папка с шаблонами")
        if not self.output_dir: errors.append("Не выбрана папка для сохранения")
        
        if errors:
            messagebox.showerror("Ошибка", "\n".join(errors))
            return
        
        required = ["01_Анкета_участника_шаблон.docx", 
                   "02_Заявка_на_участие_в_закупке_шаблон.docx",
                   "03_Предложение_о_цене_договора_шаблон.docx"]
        missing = [f for f in required if not (self.template_dir / f).exists()]
        
        if missing:
            messagebox.showerror("Ошибка", "Не найдены шаблоны:\n" + "\n".join(missing))
            return
        
        self.progress_frame.pack(fill="x", pady=(0, 8))
        self.gen_btn.configure(state="disabled", fg_color=IOS_GRAY)
        
        thread = threading.Thread(target=self._do_generate, args=(profile, tender, calc))
        thread.daemon = True
        thread.start()
    
    def _do_generate(self, profile_path, tender_path, calc_path):
        try:
            with open(profile_path, 'r', encoding='utf-8') as f:
                profile = json.load(f)
            with open(tender_path, 'r', encoding='utf-8') as f:
                tender = json.load(f)
            with open(calc_path, 'r', encoding='utf-8') as f:
                calc = json.load(f)
            
            generator = DocumentGenerator(self.template_dir, self.output_dir, self._progress_callback)
            results = generator.generate_all(profile, tender, calc)
            
            success = sum(1 for r in results if r['status'] == 'success')
            errors = [r.get('error', '') for r in results if r['status'] == 'error']
            
            if errors:
                self.root.after(0, lambda: self._show_error("\n".join(errors)))
            else:
                self.root.after(0, lambda: self._show_success(success))
        except Exception as e:
            self.root.after(0, lambda: self._show_error(str(e)))
    
    def _show_success(self, count):
        self.status_label.configure(text=f"✓ Успешно! {count} документов", text_color=IOS_GREEN)
        self.progress_bar.set(1.0)
        messagebox.showinfo("Готово!", f"Сгенерировано {count} документов!\n\nСохранено в:\n{self.output_dir}")
        self.gen_btn.configure(state="normal", fg_color=IOS_GREEN)
    
    def _show_error(self, error):
        self.status_label.configure(text="✗ Ошибка генерации", text_color=IOS_RED)
        messagebox.showerror("Ошибка", f"Произошла ошибка:\n{error}")
        self.gen_btn.configure(state="normal", fg_color=IOS_GREEN)
    
    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = TenderGeneratorApp()
    app.run()
